import docx
import sys

def extract_docx(path, output_path):
    doc = docx.Document(path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=== PARAGRAPHS ===\n")
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(p.text + "\n")
        
        f.write("\n=== TABLES ===\n")
        for table in doc.tables:
            for row in table.rows:
                # Merge cells can cause issues, but this is a simple extraction
                row_text = []
                for cell in row.cells:
                    text = cell.text.strip().replace("\n", " ")
                    if not row_text or text != row_text[-1]: # Avoid duplicate text from merged cells
                        row_text.append(text)
                f.write(" | ".join(row_text) + "\n")
            f.write("-" * 20 + "\n")

if __name__ == "__main__":
    extract_docx(sys.argv[1], sys.argv[2])
