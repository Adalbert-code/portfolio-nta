import docx
import sys

def extract_docx(path):
    doc = docx.Document(path)
    print("=== PARAGRAPHS ===")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    
    print("\n=== TABLES ===")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_text))
        print("-" * 20)

if __name__ == "__main__":
    extract_docx(sys.argv[1])
